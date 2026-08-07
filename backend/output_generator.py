"""
output_generator.py
Generates Excel/GUI output using the standardized Test Case Template
column structure (19 mandatory columns, fixed order):

   1  Requirement_ID
   2  TC_ID                    -- merged test_case_id + scenario_id
   3  Test Objective
   4  Test Details Description
   5  Test Pre-Condition
   6  Inputs                   -- single consolidated field: "Input1=Value1; Input2=Value2"
   7  Test Steps
   8  Requirement_Type         -- Functional / Non-Functional
   9  Coverage_Type            -- Validation / Verification / Integration
  10  Test_Level               -- Unit / Integration / System
  11  Scenario_Type            -- passed through from the underlying generator
  12  Expected Outputs         -- single consolidated field: "Output1=Value1; Output2=Value2"
  13  Module
  14  Safety_Level             -- High / Low (domain default, editable)
  15  Priority                 -- P1 / P2 / P3
  16  PASS/FAIL Criteria
  17  Configure_Baseline       -- kept blank
  18  Remarks
  19  Standard_Reference       -- domain document name (domain default, editable)

No more dynamic per-signal sub-columns: Inputs and Expected Outputs are each
a single column now (see _consolidated_inputs / _consolidated_outputs).
"""

import io
import re
import hashlib
from datetime import datetime
from typing import List, Dict, Tuple, Optional

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import TestCase
from config import ENGINE
from constants import DOMAIN_DEFAULTS, TEST_LEVEL_INTEGRATION_OVERRIDE


# ─── STYLING — Uniform colour scheme (Requirement 6) ─────────────────────────
# All header cells use the same blue fill; no per-column different colours.
HEADER_FILL   = PatternFill("solid", fgColor="4472C4")   # uniform blue
HEADER_FONT   = Font(bold=True, color="FFFFFF", size=10, name="Calibri")
HEADER_ALIGN  = Alignment(horizontal="center", vertical="center", wrap_text=True)

SUBHDR_FILL   = PatternFill("solid", fgColor="4472C4")   # same blue for sub-headers (Req 6)
SUBHDR_FONT   = Font(bold=True, color="FFFFFF", size=9,  name="Calibri")
SUBHDR_ALIGN  = Alignment(horizontal="center", vertical="center", wrap_text=True)

BODY_FONT     = Font(size=9, name="Calibri")
BODY_ALIGN    = Alignment(vertical="top", wrap_text=True)
CENTER_ALIGN  = Alignment(horizontal="center", vertical="top", wrap_text=True)

THIN_SIDE     = Side(style="thin", color="CCCCCC")
THIN_BORDER   = Border(left=THIN_SIDE, right=THIN_SIDE, top=THIN_SIDE, bottom=THIN_SIDE)

ALT_FILL      = PatternFill("solid", fgColor="EEF2F9")   # alternating row shading


# ─── SIGNAL EXTRACTION ────────────────────────────────────────────────────────

# Accept both "Signal: Value" (rule-based) and "Signal = Value" (Claude AI) formats
_KV_COLON = re.compile(r'^(.+?):\s*(.+)$')
_KV_EQUAL  = re.compile(r'^(.+?)\s*=\s*(.+)$')

# TC methodologies that always produce proper "SignalName: Value" inputs
# Phrases that disqualify an output signal name
_OUTPUT_SKIP_PHRASES = {
    "system successfully", "response is", "data is", "result is",
    "all sub", "no data", "logic module", "specification",
    "test case", "is correct", "the logic", "and sets",
    "for scenario", "scenario sc", "this single",
    "the output", "all conditions", "sub-requirements",
    "no gaps", "the system", "collectively", "correctly",
    "evaluates to", "independence criterion", "causes the",
    # Generic placeholders that Claude AI or the fallback path may produce —
    # these are never real signal names and must not appear as column sub-headers.
    "output signal", "output value", "output state", "expected output",
    "signal output", "signal value", "signal name",
}
_OUTPUT_STARTER_SKIP = {
    "the", "a", "an", "this", "all", "no", "for",
    "system", "response", "data", "result", "and",
    "output",   # prevents "output signal", "output value" etc. from becoming a signal name
}


_TRANSITION_PHRASE = re.compile(
    r'(?:transition(?:s|ing)?\s+from\s+)?'
    r'([A-Za-z][\w]{0,20}(?:\s+[A-Za-z][\w]{0,20}){0,2})'
    r'\s+to\s+'
    r'([A-Za-z][\w]{0,20}(?:\s+[A-Za-z][\w]{0,20}){0,2})'
    r'\s*$',
    re.IGNORECASE,
)
# Left-side words that mean "assign to X", not "from state X" — guards
# against misreading values like "Set to True" as a state transition.
_TRANSITION_LEFT_STOPWORDS = {
    "set", "reset", "change", "changed", "switch", "switched",
    "move", "moved", "go", "goes", "went", "assign", "assigned",
}


def _normalize_transition_value(value: str) -> str:
    """
    Req 2: represent state transitions with arrow notation
    ('Invalid -> Valid') instead of prose ('transition from Invalid to
    Valid'). Only rewrites values that name two short discrete states —
    a numeric phrase like '10 to 50' is left untouched because both
    captured groups must start with a letter, and assignment phrasing
    like 'Set to True' is left untouched too.
    """
    if not value or " to " not in value.lower():
        return value
    m = _TRANSITION_PHRASE.search(value.strip())
    if not m:
        return value
    left, right = m.group(1).strip(), m.group(2).strip()
    if left.split()[-1].lower() in _TRANSITION_LEFT_STOPWORDS:
        return value
    return f"{left} -> {right}"


def _parse_signal_value(entry: str) -> Tuple[str, str]:
    """Parses 'Name: Value' or 'Name = Value' into (name, value)."""
    s = entry.strip()
    # Prefer colon separator (standard format), fall back to equals
    m = _KV_COLON.match(s) or _KV_EQUAL.match(s)
    if m:
        return m.group(1).strip(), _normalize_transition_value(m.group(2).strip())
    return s, ""


def _has_numeric_inputs(tc: TestCase) -> bool:
    """True if at least one input value is numeric — used to decide whether
    min/max language is applicable (Req 5: it isn't, for pure Boolean/Enum
    inputs)."""
    for entry in tc.inputs:
        _, value = _parse_signal_value(entry)
        v = (value or "").strip().lstrip("+-")
        if v.replace(".", "", 1).isdigit():
            return True
    return False


def _extract_all_output_signals_with_values(
    expected_outcome: str,
    exclude_names: set = None,
) -> List[Tuple[str, str]]:
    """
    Returns list of (signal_name, value) pairs from expected_outcome.

    Handles ALL formats:
      "SignalName = Value. ..."                       <- rule-based standard
      "For scenario SC_N: Signal = V; Signal2 = V2." <- decision table
      " = Value. AND-decision evaluates..."           <- empty output_name
      "...sets SignalName to Value..."                <- Claude AI prose
      "...SignalName is Value..."                     <- Claude AI prose

    If exclude_names is provided (set of input signal names),
    those names are excluded from output detection.
    """
    if not expected_outcome:
        return []

    excl = {n.lower() for n in (exclude_names or set())}
    results: List[Tuple[str, str]] = []
    seen: set = set()

    def _add(name: str, value: str) -> bool:
        """Validates and adds (name, value) if not already seen."""
        name  = re.sub(r'^(?:the|a|an|this)\s+', '', name, flags=re.IGNORECASE).strip()
        if not name or len(name) < 2:
            return False
        words = name.split()
        if not (1 <= len(words) <= 8):
            return False
        if words[0].lower() in _OUTPUT_STARTER_SKIP:
            return False
        if any(p in name.lower() for p in _OUTPUT_SKIP_PHRASES):
            return False
        if name.lower() in excl:
            return False
        if name.lower() in seen:
            return False
        seen.add(name.lower())
        results.append((name, value))
        return True

    # ── Strategy 1: first clause "Signal = Value" (standard / decision table) ─
    if "=" in expected_outcome:
        first_clause = expected_outcome.split(".")[0].strip()
        if ":" in first_clause:
            after = first_clause.rsplit(":", 1)[1].strip()
            if "=" in after:
                first_clause = after
        if "=" in first_clause:
            for part in first_clause.split(";"):
                part = part.strip()
                if "=" not in part:
                    continue
                pts = part.split("=", 1)
                raw_name  = pts[0].strip()
                value_raw = pts[1].strip() if len(pts) > 1 else ""
                name  = re.sub(r'^(?:the|a|an|this)\s+', '', raw_name, flags=re.IGNORECASE).strip()
                value = value_raw.split()[0].rstrip(".,;:") if value_raw else ""
                if name and value:
                    _add(name, value)

    # ── Strategy 2: full scan for "Signal = BoolValue" (handles prose) ────────
    # Used when strategy 1 finds nothing or for additional signals
    _BOOL_VALS = r'(True|False|TRUE|FALSE|Enable|Disable|Active|Inactive|Enabled|Disabled|1|0)'
    for m in re.finditer(
        r'([A-Z][\w\s]{2,60}?)\s*=\s*' + _BOOL_VALS + r'\b',
        expected_outcome
    ):
        cand = m.group(1).strip()
        val  = m.group(2).strip()
        _add(cand, val)

    # ── Strategy 3: "sets SignalName to Value" ────────────────────────────────
    # e.g. "...sets Altitude Alert Condition Enabled to True..."
    for m in re.finditer(
        r'sets?\s+([A-Z][\w\s]{2,50}?)\s+to\s+[\'"]?' + _BOOL_VALS + r'[\'"]?\b',
        expected_outcome, re.IGNORECASE
    ):
        _add(m.group(1).strip(), m.group(2).strip())

    # ── Strategy 4: "SignalName [output] is set to 'Value'" ──────────────────
    # Catches Claude AI pattern: "Altitude Alert Condition Enabled output is set to 'True'"
    # or: "Is Enabled of Auto Start State is set to 'Enable'"
    for m in re.finditer(
        r'([A-Z][\w\s]{2,60}?)\s+(?:output\s+)?is\s+set\s+to\s+[\'"]?' + _BOOL_VALS + r'[\'"]?',
        expected_outcome, re.IGNORECASE
    ):
        cand = re.sub(r'\s+output\s*$', '', m.group(1).strip(), flags=re.IGNORECASE).strip()
        _add(cand, m.group(2).strip())

    # ── Strategy 5: "SignalName is/equals/becomes Value" ──────────────────────
    # e.g. "Altitude Alert Condition Enabled is True"
    for m in re.finditer(
        r'([A-Z][\w\s]{2,50}?)\s+(?:is|equals?|becomes?)\s+[\'"]?' + _BOOL_VALS + r'[\'"]?\b',
        expected_outcome
    ):
        _add(m.group(1).strip(), m.group(2).strip())

    return results


def _extract_output_signal(expected_outcome: str) -> Tuple[str, str]:
    """
    Extracts (signal_name, value) from expected_outcome.

    Handles all formats:
      "SignalName = Value. ..."
      "For scenario SC_N: SignalName = Value. ..."
      " = Value. AND-decision..."  -- empty output_name fallback: scans full text
    """
    if not expected_outcome or '=' not in expected_outcome:
        return "", ""

    first_clause = expected_outcome.split('.')[0].strip()

    # Strip "For scenario SC_N:" prefix
    if ':' in first_clause:
        after = first_clause.rsplit(':', 1)[1].strip()
        if '=' in after:
            first_clause = after

    if '=' not in first_clause:
        return "", ""

    parts = first_clause.split('=', 1)
    if len(parts) != 2:
        return "", ""

    raw_name  = parts[0].strip()
    value_raw = parts[1].strip()
    value     = value_raw.split()[0].rstrip('.,;:') if value_raw else ""
    if not value:
        return "", ""

    name = re.sub(r'^(?:the|a|an|this)\s+', '', raw_name, flags=re.IGNORECASE).strip()

    # Empty signal name (output_name was empty in generator)
    # Scan full expected_outcome for "SignalName = BoolValue" pattern
    if not name or len(name) < 2:
        for m in re.finditer(
            r'([A-Z][\w\s]{2,50}?)\s*=\s*(True|False|TRUE|FALSE|Active|Inactive|1|0)\b',
            expected_outcome
        ):
            candidate  = m.group(1).strip()
            cand_val   = m.group(2).strip()
            cand_words = candidate.split()
            if not (1 <= len(cand_words) <= 8):
                continue
            if cand_words[0].lower() in _OUTPUT_STARTER_SKIP:
                continue
            if any(p in candidate.lower() for p in _OUTPUT_SKIP_PHRASES):
                continue
            return candidate, cand_val
        return "", ""

    # Validate signal name
    words = name.split()
    if not (1 <= len(words) <= 8):
        return "", ""
    if words[0].lower() in _OUTPUT_STARTER_SKIP:
        return "", ""
    if any(phrase in name.lower() for phrase in _OUTPUT_SKIP_PHRASES):
        return "", ""

    return name, value


def _normalise_signal_name(name: str) -> str:
    """
    Returns a canonical form of a signal name for deduplication purposes.

    Collapses whitespace, lowercases, and strips common scenario-type
    qualifiers that Claude AI sometimes appends to signal names when
    generating multiple scenario TCs (normal / boundary / edge / robustness).

    Examples:
      "Tail Low Condition"          → "tail low condition"
      "tail low condition"          → "tail low condition"
      "TAIL LOW CONDITION"          → "tail low condition"
      "Tail Low Condition (normal)" → "tail low condition"
      "Tail Low Condition_boundary" → "tail low condition"
      "Tail Low Condition - edge"   → "tail low condition"
    """
    s = re.sub(r'\s+', ' ', name.strip()).lower()
    # Strip trailing scenario-type qualifiers added by Claude AI
    _QUALIFIERS = (
        r'\s*[\(\[]\s*(?:normal|boundary|edge|robustness|positive|negative|'
        r'baseline|flip|invalid|valid|min|max|minimum|maximum)\s*[\)\]]',
        r'\s*[-_]\s*(?:normal|boundary|edge|robustness|positive|negative|'
        r'baseline|flip|invalid|valid|min|max|minimum|maximum)\s*$',
    )
    for pat in _QUALIFIERS:
        s = re.sub(pat, '', s, flags=re.IGNORECASE).strip()
    return s


# ─── FIELD HELPERS ────────────────────────────────────────────────────────────

def _list_to_str(value) -> str:
    if isinstance(value, list):
        return "\n".join(str(v) for v in value if v)
    return str(value) if value else ""


def _module_alpha_only(module: str) -> str:
    """Requirement 7: keep only alphabetical characters and spaces."""
    cleaned = re.sub(r'[^A-Za-z\s]', '', module).strip()
    return re.sub(r'\s+', ' ', cleaned) or "General"


# ─── REQUIREMENT 5: Column F content ─────────────────────────────────────────
# Col F = Test Precondition, but per Req 5 it must consolidate:
#   • Test Objective (from col D/E)
#   • Test Steps that are related to the identified input parameter names (from H, I... cols)

def _varied_signal_names(tc: TestCase, siblings: List[TestCase], limit: int) -> str:
    """
    For an MC/DC independence-pair row (boundary scenario, SC_002+), finds the
    signal name(s) that actually differ from this TC's own SC_001/normal
    baseline sibling. Returns "" if no baseline is found or nothing differs,
    so the caller can fall back to the generic first-N-signals behaviour.
    """
    baseline = next(
        (s for s in siblings
         if s.test_case_id == tc.test_case_id
         and s.scenario_id != tc.scenario_id
         and (s.scenario_type or "").lower() == "normal"),
        None,
    )
    if not baseline:
        return ""

    skip = ("test environment", "all prerequisite", "sub-requirements")
    base_map = {
        name: val for name, val in (_parse_signal_value(e) for e in baseline.inputs)
        if name and name.lower() not in skip
    }
    tc_map = {
        name: val for name, val in (_parse_signal_value(e) for e in tc.inputs)
        if name and name.lower() not in skip
    }

    varied = [name for name, val in tc_map.items()
              if name in base_map and base_map[name] != val]
    if not varied:
        return ""
    return " and ".join(varied[:limit])


def _description_signal_names(tc: TestCase, limit: int = 2,
                              siblings: Optional[List[TestCase]] = None) -> str:
    """Pulls 1-2 real input signal names from this TC for a context-specific
    description, instead of a generic phrase.

    For MC/DC independence-pair rows this used to always name the first N
    signals declared on the TC regardless of which one this specific
    scenario actually isolates — so every boundary row under the same
    TC_ID quoted the same one or two signals even when a different signal
    was the one being flipped. When `siblings` (the other scenarios sharing
    this TC_ID) is supplied, prefer the signal(s) that differ from the
    SC_001/normal baseline — that is the signal genuinely under test here.
    """
    if siblings:
        varied = _varied_signal_names(tc, siblings, limit)
        if varied:
            return varied

    names = []
    for entry in tc.inputs:
        name, _ = _parse_signal_value(entry)
        if name and name.lower() not in ("test environment", "all prerequisite", "sub-requirements"):
            names.append(name)
        if len(names) >= limit:
            break
    return " and ".join(names) if names else "the input conditions"


def _description_output_name(tc: TestCase) -> str:
    """Pulls the real output signal name from expected_outcome when available."""
    name, _ = _extract_output_signal(tc.expected_outcome)
    return name if name else "the system output"


def _description_variant_index(tc: TestCase, modulo: int) -> int:
    """
    Deterministic (not random) variant picker: the same TC always renders
    the same phrasing on regeneration, but different requirements/scenarios
    land on different variants — avoiding repetitive, template-identical
    descriptions across the sheet (Req 3).
    """
    if modulo <= 1:
        return 0
    key = f"{tc.traceability_req_id}|{tc.scenario_id}|{tc.scenario_type}"
    return int(hashlib.md5(key.encode()).hexdigest(), 16) % modulo


def _col_e_test_details(tc: TestCase, siblings: Optional[List[TestCase]] = None) -> str:
    """
    Column E — Test Details Description.
    A narrative, requirement-specific explanation of what this test does and
    why — scenario-type driven, but phrased with the real signal/output names
    and varied sentence structure so it isn't identical across every TC of
    the same scenario type (Req 3).
    Does NOT repeat inputs, expected output, preconditions, or the test
    objective, and does NOT restate Design Methodology / Module, since those
    already have their own dedicated columns (Req 1).

    `siblings` — the other scenarios sharing this TC_ID — lets MC/DC
    (boundary) rows name the signal actually being isolated in THIS row
    rather than always naming the first declared input(s); see
    _description_signal_names.
    """
    sig = _description_signal_names(tc, siblings=siblings)
    out = _description_output_name(tc)
    req = tc.traceability_req_id or "the requirement"

    _DETAIL_VARIANTS = {
        "normal": [
            f"Sets {sig} to their nominal, required values simultaneously and confirms "
            f"{out} activates as {req} specifies. Also serves as the MC/DC baseline for "
            f"the independence-pair tests derived from this decision.",
            f"Drives the primary activation path for {req}: every condition at its correct "
            f"operating value at once, checking that {out} responds exactly as intended.",
            f"Establishes the reference 'all-conditions-met' case for {req}, against which "
            f"the MC/DC flip tests for {sig} are later compared.",
        ],
        "boundary": [
            f"Flips {sig} on its own while holding every other condition at its required "
            f"value, confirming {out} changes solely because of that one variable — the "
            f"MC/DC independence check for {req}.",
            f"Isolates {sig} from the other conditions in {req} to prove it independently "
            f"controls {out}, satisfying the MC/DC coverage objective for this decision.",
            f"Tests whether {out} tracks a single-condition change in {sig} without being "
            f"masked by any other input, per the MC/DC pairing required for {req}.",
        ],
        "edge": [
            f"Puts {sig} at an inactive or non-triggering value at the same time as every "
            f"other condition and confirms {out} stays safely inactive under {req}.",
            f"Checks the fully de-energised corner case for {req}: all inputs simultaneously "
            f"at non-activating values, with {out} expected to remain in its safe state.",
            f"Exercises the combined worst-case configuration where none of {sig} meet the "
            f"activation criteria, verifying {out} does not activate unintentionally.",
        ],
        "robustness": [
            f"Feeds {sig} an invalid or out-of-range value and confirms the system tolerates "
            f"it without crashing, corrupting {out}, or entering an undefined state.",
            f"Checks fault tolerance for {req} by corrupting {sig}, then verifying the system "
            f"degrades safely and {out} does not produce an undefined result.",
            f"Verifies recovery behaviour: after {sig} receives a bad value, restoring it to "
            f"a valid range should bring {out} back to its correct state.",
        ],
        "transition": [
            f"Drives {sig} through the state change called out in {req} and confirms {out} "
            f"activates or deactivates at the correct trigger point.",
            f"Walks {sig} across the transition boundary defined in {req}, checking that "
            f"{out} switches state exactly when it should — no early or delayed response.",
            f"Confirms partial or sequenced changes in {sig} are handled correctly as the "
            f"system moves between states, with {out} reflecting the right state at each step.",
        ],
        "invalid_input": [
            f"Drives {sig} below its declared ICD minimum or above its declared ICD maximum "
            f"and confirms the system detects the out-of-range value and rejects/clamps it "
            f"per specification for {req}, without crashing or corrupting {out}.",
            f"Exercises the ICD-declared invalid region for {sig}, checking that {req} is "
            f"enforced against out-of-declared-range inputs rather than silently accepting them.",
        ],
    }

    sc_type  = (tc.scenario_type or "").lower()
    variants = _DETAIL_VARIANTS.get(sc_type, [f"Verifies the functional behaviour of {req} as specified."])
    idx      = _description_variant_index(tc, len(variants))
    return variants[idx]

def _col_f_precondition(tc: TestCase, input_signals: List[str]) -> str:
    """
    Column F — Test Precondition.
    Contains ONLY the actual preconditions for the test case.
    Does NOT include test objective, test steps, or pre-set input values
    (those belong in their own dedicated columns).
    """
    if not tc.preconditions:
        return ""
    return _list_to_str(tc.preconditions)


# ─── REQUIREMENT 8: Remarks bullet formatting ─────────────────────────────────

def _remarks_bullets(tc: TestCase) -> str:
    """
    Requirement 8:
    - Remove test-basis-related info
    - Testing Type / Scenario Type are NOT repeated here — they already have
      dedicated columns in the test case template, so including them in the
      Remarks/Additional Information text would be redundant.
    - Describe what is tested in each SC (e.g. INPUT_1 maximum value is tested)
    - Bullet-point format
    """
    bullets = []

    # What is being tested (Req 8 — describe each SC)
    if tc.scenario_type == "boundary" and not _has_numeric_inputs(tc):
        boundary_what = "Each declared valid/invalid state of the input parameter(s) is exercised individually (no numeric min/max applies)."
    else:
        boundary_what = "Input boundary values tested: minimum, maximum, min-1, max+1 for each parameter."
    sc_what = {
        "normal":     "All input values set to normal/valid values; correct system output is verified.",
        "boundary":   boundary_what,
        "edge":       "Edge case conditions tested (state transitions, simultaneous changes, unusual-but-valid states).",
        "robustness": "Invalid/out-of-range input values tested; system must respond safely without crash.",
        "invalid_input": "Input parameter's declared valid range (per ICD/supporting-document cross-reference) "
                          "is exercised at and beyond its boundaries, including below-minimum and above-maximum "
                          "invalid values; system must detect and reject these safely.",
    }
    bullets.append(f"• What is tested: {sc_what.get(tc.scenario_type, 'Functional system behaviour verified.')}")

    # Per-input description (e.g. "INPUT_1 maximum value is tested")
    for entry in tc.inputs:
        name, value = _parse_signal_value(entry)
        if name and value and name.lower() not in ("test environment", "all prerequisite", "sub-requirements"):
            if tc.scenario_type == "boundary":
                if "max" in value.lower() or "maximum" in value.lower():
                    bullets.append(f"• {name}: maximum value is tested")
                elif "min" in value.lower() or "minimum" in value.lower():
                    bullets.append(f"• {name}: minimum value is tested")
                elif "-1" in value or "below" in value.lower():
                    bullets.append(f"• {name}: below-minimum value is tested (invalid range)")
                elif "+1" in value or "above" in value.lower():
                    bullets.append(f"• {name}: above-maximum value is tested (invalid range)")
                else:
                    bullets.append(f"• {name}: boundary value '{value}' is tested")
            elif tc.scenario_type == "edge":
                bullets.append(f"• {name}: edge-case value '{value}' is tested (state-transition condition)")
            elif tc.scenario_type == "robustness":
                bullets.append(f"• {name}: invalid/out-of-range value '{value}' is tested")

    # Input source note (Req 4)
    inputs_raw = " ".join(tc.inputs).lower()
    if any(kw in inputs_raw for kw in ["icd", "derived", "interface"]):
        bullets.append("• Input source: Values derived from ICD document (not explicitly defined in SRS).")
    else:
        bullets.append("• Input source: Input values explicitly defined in SRS specification.")

    # Sub-requirements / cross-refs from raw remarks (strip test-basis lines)
    if tc.remarks:
        raw_parts = re.split(r'\s*[\|\n•]+\s*', tc.remarks)
        for part in raw_parts:
            part = part.strip()
            if not part:
                continue
            # Remove test-basis lines (Req 8)
            if re.search(
                r'test\s+basis|input\s+values\s+derived\s+from\s+srs|srs\s+requirement\s+\w',
                part, re.IGNORECASE
            ):
                continue
            # Include enum definitions, sub-req refs, notes
            if re.search(r'enum|sub.req|note|reference|derived from icd|document context', part, re.IGNORECASE):
                bullets.append(f"• {part}")

    return "\n".join(bullets)


# ─── REQUIREMENT 10: Depends On ───────────────────────────────────────────────

def _depends_on(raw_dep: str, tc_id: str, sc_no: int) -> str:
    """
    Depends On column.
    Format: TC_UT_001_SC-001  (hyphen between SC and number)

    The generator writes:
      - "None"            for SC_001 (baseline)
      - "TC_UT_001_SC-001" for SC_002+ (always references baseline with hyphen)

    This function passes the value through unchanged if already formatted,
    or applies a fallback for legacy/MCP data.
    """
    if not raw_dep or raw_dep.strip().lower() == "none":
        return "None"
    raw = raw_dep.strip()
    # Already formatted (TC_ID_SC-001 hyphen format or TC_ID_SC_001 underscore)
    if "_SC-" in raw or "_SC_" in raw.upper():
        return raw
    # Fallback: bare TC_ID — append SC-001 (baseline reference, hyphen format)
    return f"{raw}_SC-001"


# ─── HEADER WRITER ────────────────────────────────────────────────────────────

STANDARD_TEMPLATE_COLUMNS: List[Tuple[str, int]] = [
    ("Requirement_ID",              21),
    ("TC_ID",                       16),
    ("Test Objective",              20),
    ("Test Details Description",    22),
    ("Test Pre-Condition",          32),
    ("Inputs",                      32),
    ("Test Steps",                  30),
    ("Requirement_Type",            14),
    ("Coverage_Type",               14),
    ("Test_Level",                  12),
    ("Scenario_Type",               14),
    ("Expected Outputs",            32),
    ("Module",                       9),
    ("Safety_Level",                10),
    ("Priority",                     9),
    ("PASS/FAIL Criteria",          40),
    ("Configure_Baseline",          16),
    ("Remarks",                     32),
    ("Standard_Reference",          18),
]


def _write_headers(ws, input_signals: List[str] = None, output_signals: List[str] = None) -> Dict[str, int]:
    """
    Writes the single-row standardized template header (19 mandatory
    columns, fixed order — see STANDARD_TEMPLATE_COLUMNS). Returns a dict of
    column-name -> column-index for use when writing data.

    `input_signals`/`output_signals` are accepted for call-site compatibility
    but no longer produce sub-columns — Inputs and Expected Outputs are each
    a single consolidated column now.
    """
    col_map: Dict[str, int] = {}
    for col, (hdr, width) in enumerate(STANDARD_TEMPLATE_COLUMNS, start=1):
        c = ws.cell(row=1, column=col, value=hdr)
        c.font = HEADER_FONT; c.fill = HEADER_FILL
        c.alignment = HEADER_ALIGN; c.border = THIN_BORDER
        ws.column_dimensions[get_column_letter(col)].width = width
        col_map[hdr] = col

    ws.row_dimensions[1].height = 28
    ws.freeze_panes = "A2"
    return col_map


# ─── OUTPUT VALUE EXTRACTOR ───────────────────────────────────────────────────

def _extract_output_value_only(expected_outcome: str) -> str:
    """
    Extracts ONLY the plain value from an expected_outcome string.

    Handles these formats and returns just the value token:
      "SignalName = True. ..."           -> "True"
      "SignalName = False. ..."          -> "False"
      "For scenario SC_001: X = True."  -> "True"
      "System successfully executes..."  -> first sentence (no signal prefix found)

    The goal is to never write "SignalName = Value" into the output cell —
    only "Value" (or a short descriptive first-sentence if no signal is present).
    """
    if not expected_outcome:
        return ""

    first_clause = expected_outcome.split('.')[0].strip()

    # Strip "For scenario SC_N:" prefix
    if ':' in first_clause:
        after = first_clause.rsplit(':', 1)[1].strip()
        if '=' in after:
            first_clause = after

    if '=' in first_clause:
        parts = first_clause.split('=', 1)
        if len(parts) == 2:
            raw_value = parts[1].strip()
            value = raw_value.split()[0].rstrip('.,;:') if raw_value else ""
            # Accept only unambiguous boolean / enum tokens as the extracted value
            _KNOWN_VALUES = {
                'true', 'false', 'enabled', 'disabled',
                'active', 'inactive', '1', '0',
                'pass', 'fail', 'yes', 'no', 'set', 'reset',
                'on', 'off', 'high', 'low', 'open', 'closed',
                'valid', 'invalid',
            }
            if value and value.lower() in _KNOWN_VALUES:
                return value

    # No clean signal=value prefix found — return the first sentence as-is
    # (covers standard TCs whose outcome starts with "System successfully…")
    return first_clause if first_clause else ""


# ─── GUI/EXPORT PARITY ────────────────────────────────────────────────────────
# Single source of truth for the narrative/derived columns (Test Details
# Description, Test Precondition, Expected Outputs, Depends On, Remarks,
# Module) so the React GUI and the Excel/Word export always show identical
# text. Previously the GUI (TCTable.jsx / ResultsTable.jsx) re-implemented
# these columns from scratch in JS — a static, non-signal-aware version of
# _col_e_test_details that had drifted from this Python implementation.
# The backend now computes them once per test case and the API attaches the
# result to the payload; the frontend must read these fields directly rather
# than recomputing them.

# ─── STANDARDIZED TEST CASE TEMPLATE — Columns 2/6/12/8/9/10/11/14/16/17/19 ───
# Single source of truth for the 19-column standardized template so the GUI
# table and the Excel export never drift, same pattern as the narrative
# columns above.

def _merged_tc_id(tc: TestCase) -> str:
    """Column 2 — TC_ID: merges test_case_id + scenario_id into one field,
    e.g. 'TC_001' + 'SC_001' -> 'TC_001_SC_001'. Falls back gracefully if
    either half is missing so a partially-populated TC never renders blank."""
    tc_id  = (tc.test_case_id or "").strip()
    sc_id  = (tc.scenario_id  or "").strip()
    if tc_id and sc_id:
        return f"{tc_id}_{sc_id}" if not tc_id.endswith(sc_id) else tc_id
    return tc_id or sc_id or ""


def _consolidated_inputs(tc: TestCase) -> str:
    """Column 6 — Inputs: consolidates every named input signal into ONE
    field, 'Input1=Value1; Input2=Value2', instead of separate columns per
    signal. Entries that aren't valid Name:Value/Name=Value pairs are kept
    as-is so free-text inputs aren't silently dropped."""
    if not tc.inputs:
        return ""
    parts = []
    for entry in tc.inputs:
        name, value = _parse_signal_value(entry)
        if name and value:
            parts.append(f"{name}={value}")
        elif entry:
            parts.append(str(entry).strip())
    return "; ".join(p for p in parts if p)


def _consolidated_outputs(tc: TestCase) -> str:
    """Column 12 — Expected Outputs: consolidates every named output signal
    into ONE field, 'Output1=ExpectedValue1; Output2=ExpectedValue2', instead
    of separate columns per signal. Falls back to the first-sentence summary
    when no named signal=value pairs can be parsed out of expected_outcome."""
    input_names = set()
    for entry in tc.inputs:
        n, _ = _parse_signal_value(entry)
        if n:
            input_names.add(_normalise_signal_name(n))
    pairs = _extract_all_output_signals_with_values(tc.expected_outcome, exclude_names=input_names)
    if pairs:
        return "; ".join(f"{name}={value}" for name, value in pairs)
    fallback = _extract_output_value_only(tc.expected_outcome) or (tc.expected_outcome or "").split(".")[0].strip()
    return fallback


def _requirement_type_display(tc: TestCase) -> str:
    """Column 8 — Requirement_Type: 'Functional' / 'Non-Functional'."""
    return "Non-Functional" if (tc.requirement_type or "").lower().startswith("non") else "Functional"


def _coverage_type_display(tc: TestCase) -> str:
    """Column 9 — Coverage_Type: 'Validation' / 'Verification' / 'Integration',
    mapped 1:1 from the existing testing_type field (no new classification)."""
    mapping = {"verification": "Verification", "validation": "Validation", "integration": "Integration"}
    return mapping.get((tc.testing_type or "").lower(), "Verification")

def _test_level_display(tc: TestCase, domain: str) -> str:
    """Column 10 — Test_Level: Unit / Integration / System.
    Per-TC override (tc.test_level, settable by the MCP/Claude-AI path)
    wins; otherwise the session-level domain default applies, promoted to
    'Integration' whenever Coverage_Type is Integration."""
    if tc.test_level:
        return tc.test_level
    if _coverage_type_display(tc) == "Integration":
        return TEST_LEVEL_INTEGRATION_OVERRIDE
    return DOMAIN_DEFAULTS.get(domain, DOMAIN_DEFAULTS["general"])["test_level"]


def _scenario_type_display(tc: TestCase) -> str:
    """Column 11 — Scenario_Type. Passed through as-is (title-cased) rather
    than collapsed into a smaller enum — Transition/Fault/Timing/Invalid_Input
    scenarios keep their real classification per Hari's decision to preserve
    existing generation logic rather than lose scenario-type fidelity."""
    raw = (tc.scenario_type or "normal").replace("_", " ")
    return "".join(w.capitalize() for w in raw.split())


def _safety_level_display(tc: TestCase, domain: str) -> str:
    """Column 14 — Safety_Level: High / Low.
    Per-TC override wins (MCP/Claude-AI path, which can reason about DAL/
    ASIL from the SRS); otherwise the session-level domain default applies.
    This is a starting point for review, not an authoritative DAL/ASIL
    classification — the tool does not track per-requirement DAL/ASIL."""
    if tc.safety_level:
        return tc.safety_level
    return DOMAIN_DEFAULTS.get(domain, DOMAIN_DEFAULTS["general"])["safety_level"]


def _standard_reference_display(tc: TestCase, domain: str) -> str:
    """Column 19 — Standard_Reference: document name(s) for the domain.
    Per-TC override wins (MCP/Claude-AI path, which cites the specific
    clause via the domain pack in general-tc-skill); otherwise the
    session-level domain default document name applies."""
    if tc.standard_reference:
        return tc.standard_reference
    return DOMAIN_DEFAULTS.get(domain, DOMAIN_DEFAULTS["general"])["standard_reference"]


def _pass_fail_criteria(tc: TestCase, outputs_display: str) -> str:
    """Column 16 — PASS/FAIL Criteria, derived from the same expected-output
    text as column 12 so the two never contradict each other."""
    expected = outputs_display or "expected result as specified"
    return f"PASS: Actual result matches {expected}. FAIL: Any deviation from the expected result."


def compute_standard_template_fields(tc: TestCase, domain: str = "general") -> dict:
    """Computes the standardized 19-column template's derived/consolidated
    fields (columns not already covered by compute_gui_display_fields) as a
    flat dict, ready to attach to a TestCase for both the GUI table and the
    Excel export."""
    outputs_display = _consolidated_outputs(tc)
    return {
        "tc_id_display":              _merged_tc_id(tc),
        "inputs_display":             _consolidated_inputs(tc),
        "outputs_display":            outputs_display,
        "requirement_type_display":   _requirement_type_display(tc),
        "coverage_type_display":      _coverage_type_display(tc),
        "test_level_display":         _test_level_display(tc, domain),
        "scenario_type_display":      _scenario_type_display(tc),
        "safety_level_display":       _safety_level_display(tc, domain),
        "pass_fail_criteria":         _pass_fail_criteria(tc, outputs_display),
        "configure_baseline":         "",   # Req: keep blank for now
        "standard_reference_display": _standard_reference_display(tc, domain),
    }


def compute_gui_display_fields(tc: TestCase, siblings: Optional[List[TestCase]] = None,
                                domain: str = "general") -> dict:
    """Computes the same derived columns used in generate_excel/generate_docx
    and returns them as a flat dict of field_name -> display string, ready to
    be attached to the TestCase (or the raw dict backing it) before it is
    sent to the GUI.

    `siblings` should be every other TestCase sharing this batch (ideally at
    least everything sharing this TC_ID) so that MC/DC boundary rows can
    identify the actually-varied signal — see _description_signal_names.

    `domain` selects the Safety_Level / Test_Level / Standard_Reference
    defaults from constants.DOMAIN_DEFAULTS (see GenerateRequest.domain).
    """
    sc_lbl = (tc.scenario_id or "").strip()
    try:
        sc_no = int(sc_lbl.replace("SC_", "")) if sc_lbl.startswith("SC_") else 1
    except ValueError:
        sc_no = 1

    expected_first_sentence = _extract_output_value_only(tc.expected_outcome)
    if not expected_first_sentence and tc.expected_outcome:
        expected_first_sentence = tc.expected_outcome.split(".")[0].strip()

    fields = {
        "test_details_description":  _col_e_test_details(tc, siblings=siblings),
        "test_precondition_display": _col_f_precondition(tc, []),
        "expected_outputs_display":  expected_first_sentence,
        "depends_on_display":        _depends_on(tc.dependent_test_cases, tc.test_case_id, sc_no),
        "remarks_display":           _remarks_bullets(tc),
        "module_display":            _module_alpha_only(tc.module),
    }
    fields.update(compute_standard_template_fields(tc, domain=domain))
    return fields


# ─── STANDALONE ROW WRITER ────────────────────────────────────────────────────

def _write_tc_row(ws, row_idx: int, tc: TestCase,
                  col_map: dict, in_sigs: List[str] = None, out_sigs: List[str] = None,
                  signal_defaults: dict = None, siblings: Optional[List[TestCase]] = None,
                  domain: str = "general") -> None:
    """
    Writes one TC row into worksheet ws at row_idx using the standardized
    19-column template. Inputs and Expected Outputs are each written as a
    single consolidated "Name1=Value1; Name2=Value2" cell (no per-signal
    sub-columns) per the standardized template spec.

    `in_sigs`/`out_sigs`/`signal_defaults` are accepted for call-site
    compatibility but no longer drive column layout.
    """
    is_alt = (row_idx % 2 == 0)

    def _p(col: int, value, center: bool = False):
        cell = ws.cell(row=row_idx, column=col, value=value)
        cell.font      = BODY_FONT
        cell.alignment = CENTER_ALIGN if center else BODY_ALIGN
        cell.border    = THIN_BORDER
        if is_alt:
            cell.fill  = ALT_FILL

    fields = compute_standard_template_fields(tc, domain=domain)
    test_details = _col_e_test_details(tc, siblings=siblings)
    precondition = _col_f_precondition(tc, [])
    remarks      = _remarks_bullets(tc)

    _p(col_map["Requirement_ID"],           tc.traceability_req_id)
    _p(col_map["TC_ID"],                    fields["tc_id_display"])
    _p(col_map["Test Objective"],           tc.objective)
    _p(col_map["Test Details Description"], test_details)
    _p(col_map["Test Pre-Condition"],       precondition)
    _p(col_map["Inputs"],                   fields["inputs_display"] or "N/A — see Test Steps")
    _p(col_map["Test Steps"],               _list_to_str(tc.test_steps))
    _p(col_map["Requirement_Type"],         fields["requirement_type_display"], center=True)
    _p(col_map["Coverage_Type"],            fields["coverage_type_display"], center=True)
    _p(col_map["Test_Level"],               fields["test_level_display"], center=True)
    _p(col_map["Scenario_Type"],            fields["scenario_type_display"], center=True)
    _p(col_map["Expected Outputs"],         fields["outputs_display"] or "N/A")
    _p(col_map["Module"],                   _module_alpha_only(tc.module), center=True)
    _p(col_map["Safety_Level"],             fields["safety_level_display"], center=True)
    _p(col_map["Priority"],                 tc.priority, center=True)
    _p(col_map["PASS/FAIL Criteria"],       fields["pass_fail_criteria"])
    _p(col_map["Configure_Baseline"],       fields["configure_baseline"])
    _p(col_map["Remarks"],                  remarks)
    _p(col_map["Standard_Reference"],       fields["standard_reference_display"], center=True)


# ─── SAFE SHEET NAME ───────────────────────────────────────────────────────────

def _safe_sheet_name(req_id: str, used: set) -> str:
    """Converts req_id to valid Excel sheet name; resolves collisions."""
    clean = re.sub(r'[\\/*?:\[\]]', '_', req_id)
    clean = re.sub(r'[,\s]+', '_', clean)
    clean = re.sub(r'_+', '_', clean).strip('_')
    base  = clean[:31]
    name  = base
    n     = 1
    while name in used:
        suffix = f"_{n:02d}"
        name   = base[:31 - len(suffix)] + suffix
        n     += 1
    return name


# ─── EXCEL EXPORT ─────────────────────────────────────────────────────────────

def generate_excel(test_cases: List[TestCase], removed_count: int, domain: str = "general") -> bytes:
    """
    Generate Excel using the standardized 19-column Test Case Template
    (see module docstring / STANDARD_TEMPLATE_COLUMNS). Each requirement
    still gets its own sheet; `domain` selects the Safety_Level/Test_Level/
    Standard_Reference defaults (see constants.DOMAIN_DEFAULTS) for any TC
    that doesn't carry a per-TC override.
    """
    wb = openpyxl.Workbook()
    # Remove the default empty sheet created by openpyxl — we do NOT want a
    # combined "test_cases" sheet; each requirement gets its own sheet instead.
    default_ws = wb.active
    wb.remove(default_ws)

    # ── Summary sheet ─────────────────────────────────────────────────────────
    ws2 = wb.create_sheet(title="Summary")
    ws2.column_dimensions["A"].width = 35
    ws2.column_dimensions["B"].width = 25

    sum_hdr_font = Font(bold=True, color="FFFFFF", size=11, name="Calibri")
    sum_hdr_fill = PatternFill("solid", fgColor="2F4F8F")
    lbl_font     = Font(bold=True, size=10, name="Calibri")
    val_font     = Font(size=10, name="Calibri")

    def _sh_title(r, text):
        c = ws2.cell(row=r, column=1, value=text)
        c.font = sum_hdr_font; c.fill = sum_hdr_fill
        ws2.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
        c.alignment = Alignment(horizontal="center")

    def _sh_row(r, label, value):
        ws2.cell(row=r, column=1, value=label).font = lbl_font
        ws2.cell(row=r, column=2, value=value).font = val_font

    from collections import Counter
    r = 1
    _sh_title(r, "Test Case Generation Summary"); r += 1
    _sh_row(r, "Total Test Cases", len(test_cases)); r += 1
    _sh_row(r, "Duplicates Removed", removed_count); r += 1
    _sh_row(r, "Generated On", datetime.now().strftime("%Y-%m-%d %H:%M:%S")); r += 2

    _sh_title(r, "By Module"); r += 1
    for mod, cnt in sorted(Counter(_module_alpha_only(tc.module) for tc in test_cases).items()):
        _sh_row(r, mod, cnt); r += 1
    r += 1

    _sh_title(r, "By Scenario Type"); r += 1
    for st, cnt in sorted(Counter(tc.scenario_type for tc in test_cases).items()):
        _sh_row(r, st.capitalize(), cnt); r += 1
    r += 1

    _sh_title(r, "By Testing Type"); r += 1
    for tt, cnt in sorted(Counter(tc.testing_type for tc in test_cases).items()):
        _sh_row(r, tt.capitalize(), cnt); r += 1

    # ── Per-requirement sheets ────────────────────────────────────────────────
    # Every unique traceability_req_id gets its own sheet with ONLY its own
    # signal columns. Requirements differing only in ID number get separate sheets.
    from collections import OrderedDict
    req_groups: OrderedDict = OrderedDict()
    for tc in test_cases:
        rid = tc.traceability_req_id
        if rid not in req_groups:
            req_groups[rid] = []
        req_groups[rid].append(tc)

    used_names: set = {ws.title for ws in wb.worksheets}

    for req_id, req_tcs in req_groups.items():
        sname    = _safe_sheet_name(req_id, used_names)
        used_names.add(sname)
        ws_r     = wb.create_sheet(title=sname)
        r_cmap   = _write_headers(ws_r)

        for row_idx, tc in enumerate(req_tcs, start=2):
            _write_tc_row(ws_r, row_idx, tc, r_cmap,
                          siblings=req_tcs, domain=domain)

    # Summary sheet always last
    wb.move_sheet("Summary", offset=len(wb.worksheets) - 1)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ─── WORD EXPORT ──────────────────────────────────────────────────────────────

def generate_docx(test_cases: List[TestCase], removed_count: int, domain: str = "general") -> bytes:
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.8)
        section.left_margin = section.right_margin = Inches(0.9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Test Case Report")
    run.font.size = Pt(20); run.font.bold = True
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        f"Total: {len(test_cases)} test cases  |  Duplicates removed: {removed_count}"
    ).font.size = Pt(9)
    doc.add_paragraph()

    from collections import defaultdict
    by_module = defaultdict(list)
    for tc in test_cases:
        by_module[_module_alpha_only(tc.module)].append(tc)

    for module in sorted(by_module.keys()):
        h = doc.add_paragraph(f"Module: {module}")
        h.style = "Heading 1"

        for tc in by_module[module]:
            fields = compute_standard_template_fields(tc, domain=domain)
            req_id = tc.traceability_req_id

            sub = doc.add_paragraph(f"{fields['tc_id_display']} | {fields['scenario_type_display']}")
            sub.style = "Heading 2"

            rows = [
                ("Requirement_ID",          req_id),
                ("TC_ID",                   fields["tc_id_display"]),
                ("Test Objective",          tc.objective),
                ("Test Details Description", _col_e_test_details(tc, siblings=by_module[module])),
                ("Test Pre-Condition",      _col_f_precondition(tc, [])),
                ("Inputs",                  fields["inputs_display"]),
                ("Test Steps",              _list_to_str(tc.test_steps)),
                ("Requirement_Type",        fields["requirement_type_display"]),
                ("Coverage_Type",           fields["coverage_type_display"]),
                ("Test_Level",              fields["test_level_display"]),
                ("Scenario_Type",           fields["scenario_type_display"]),
                ("Expected Outputs",        fields["outputs_display"]),
                ("Module",                  _module_alpha_only(tc.module)),
                ("Safety_Level",            fields["safety_level_display"]),
                ("Priority",                tc.priority),
                ("PASS/FAIL Criteria",      fields["pass_fail_criteria"]),
                ("Configure_Baseline",      fields["configure_baseline"]),
                ("Remarks",                 _remarks_bullets(tc)),
                ("Standard_Reference",      fields["standard_reference_display"]),
            ]

            table = doc.add_table(rows=len(rows), cols=2)
            table.style = "Table Grid"
            for ri, (label, val) in enumerate(rows):
                row = table.rows[ri]
                lc = row.cells[0]; lc.width = Inches(2.0)
                lr = lc.paragraphs[0].add_run(label)
                lr.font.bold = True; lr.font.size = Pt(9)
                vr = row.cells[1].paragraphs[0].add_run(str(val))
                vr.font.size = Pt(9)

            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()